import re
import glob
from docx import Document
def task_func(directory_path: str) -> int:
    # Initialize a counter for the number of files processed
    files_processed = 0
    
    # Use glob to find all .docx files in the specified directory
    docx_files = glob.glob(f"{directory_path}/*.docx")
    
    # Iterate over each .docx file
    for file_path in docx_files:
        # Open the document
        doc = Document(file_path)
        
        # Iterate over each paragraph in the document
        for paragraph in doc.paragraphs:
            # Use regular expression to find all double quotes and add a backslash before each
            paragraph.text = re.sub(r'"', r'\"', paragraph.text)
        
        # Iterate over each table in the document
        for table in doc.tables:
            # Iterate over each cell in the table
            for row in table.rows:
                for cell in row.cells:
                    # Use regular expression to find all double quotes and add a backslash before each
                    cell.text = re.sub(r'"', r'\"', cell.text)
        
        # Save the modified document
        doc.save(file_path)
        
        # Increment the counter
        files_processed += 1
    
    # Return the number of files processed
    return files_processed