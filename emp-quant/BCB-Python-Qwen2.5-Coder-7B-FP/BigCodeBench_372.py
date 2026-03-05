import re
import glob
from docx import Document
def task_func(directory_path: str) -> int:
    # Initialize the count of processed files
    processed_files = 0
    
    # Use glob to find all .docx files in the provided directory
    docx_files = glob.glob(directory_path + '/*.docx')
    
    # Loop through each .docx file
    for file_path in docx_files:
        # Load the document
        doc = Document(file_path)
        
        # Loop through each paragraph in the document
        for paragraph in doc.paragraphs:
            # Use regular expression to find double quotes and add a backslash before each
            paragraph.text = re.sub(r'"', r'\\"', paragraph.text)
        
        # Loop through each table in the document
        for table in doc.tables:
            # Loop through each cell in the table
            for row in table.rows:
                for cell in row.cells:
                    # Use regular expression to find double quotes and add a backslash before each
                    cell.text = re.sub(r'"', r'\\"', cell.text)
        
        # Save the modified document
        doc.save(file_path)
        
        # Increment the count of processed files
        processed_files += 1
    
    # Return the number of .docx files processed
    return processed_files