
import re
import glob
from docx import Document

def task_func(directory_path: str) -> int:
    processed_files = 0
    # Iterate over all .docx files in the provided directory
    for file_path in glob.glob(directory_path + '/*.docx'):
        # Load the document
        doc = Document(file_path)
        # Iterate over all paragraphs in the document
        for paragraph in doc.paragraphs:
            # Replace double quotes with escaped double quotes
            paragraph.text = re.sub(r'"', r'\\"', paragraph.text)
        # Iterate over all tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Replace double quotes with escaped double quotes in cell text
                    cell.text = re.sub(r'"', r'\\"', cell.text)
        # Save the modified document
        doc.save(file_path)
        processed_files += 1
    return processed_files