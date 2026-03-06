import re
import glob
from docx import Document
def task_func(directory_path: str) -> int:
    """
    Processes all Word (.docx) files in the provided directory, searching for double quotes in the text 
    and adding a backslash before each double quote to "protect" it.
    
    Parameters:
    - directory_path (str): Path to the directory containing .docx files to be processed.
    
    Returns:
    - int: Number of .docx files processed.

    Requirements:
    - re
    - docx
    - glob

    Example:
    >>> import tempfile
    >>> temp_dir = tempfile.mkdtemp()
    >>> doc = Document()
    >>> _ = doc.add_paragraph("This is a sample text with double quotes.")
    >>> doc.save(temp_dir + '/sample.docx')
    >>> task_func(temp_dir)
    1
    """
    # Get all .docx files in the directory
    docx_files = glob.glob(directory_path + '/*.docx')
    # Process each .docx file
    for docx_file in docx_files:
        # Open the .docx file
        doc = Document(docx_file)
        # Get all paragraphs in the .docx file
        paragraphs = doc.paragraphs
        # Process each paragraph
        for paragraph in paragraphs:
            # Get the text in the paragraph
            text = paragraph.text
            # Search for double quotes in the text
            matches = re.findall(r'"', text)
            # If there are double quotes in the text, add a backslash before each double quote
            if matches:
                # Replace double quotes with double quotes with a backslash before them
                paragraph.text = re.sub(r'"', r'\"', text)
        # Save the .docx file
        doc.save(docx_file)
    # Return the number of .docx files processed
    return len(docx_files)