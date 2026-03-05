import re
import glob
from docx import Document
def task_func(directory_path: str) -> int:
    # Initialize a counter for processed files
    processed_files = 0
    
    # Use glob to find all .docx files in the directory
    docx_files = glob.glob(f"{directory_path}/*.docx")
    
    # Loop through each .docx file
    for file_path in docx_files:
        # Load the document
        doc = Document(file_path)
        
        # Loop through each paragraph in the document
        for paragraph in doc.paragraphs:
            # Replace double quotes with escaped double quotes
            paragraph.text = re.sub(r'"', r'\"', paragraph.text)
        
        # Loop through each table in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Replace double quotes with escaped double quotes in cell text
                    cell.text = re.sub(r'"', r'\"', cell.text)
        
        # Save the modified document
        doc.save(file_path)
        
        # Increment the counter for processed files
        processed_files += 1
    
    # Return the number of processed files
    return processed_files